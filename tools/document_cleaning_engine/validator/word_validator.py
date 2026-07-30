"""WordValidator — Word 清理结果验证器。

检查清理后的 DOCX 文件完整性：
-  可打开（python-docx）
-  ZIP 结构完整
-  Section 数量一致
-  Header/Footer 变化合理
-  文本变化在预期范围内
"""

from __future__ import annotations

import logging
import zipfile
from typing import Dict, List, Optional
from xml.etree import ElementTree as ET

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

logger = logging.getLogger(__name__)

# OOXML 必需文件
REQUIRED_XML_FILES = {
    "word/document.xml",
    "word/styles.xml",
}


class WordValidator:
    """Word 清理结果验证器。

    对清理前后的 DOCX 进行结构和内容比较验证。
    """

    def validate(
        self,
        original_path: str,
        cleaned_path: str,
    ) -> Dict[str, object]:
        """验证 Word 清理结果。

        Args:
            original_path: 原始 DOCX 路径。
            cleaned_path: 清理后 DOCX 路径。

        Returns:
            验证结果字典。
        """
        result: Dict[str, object] = {
            "open_success": False,
            "zip_structure_ok": False,
            "structure": {},
            "content": {},
        }

        # ZIP 结构检查（不依赖 python-docx）
        zip_ok = self._check_zip_structure(cleaned_path)
        result["zip_structure_ok"] = zip_ok

        if not zip_ok:
            result["structure"] = {
                "section_match": False,
                "original_sections": 0,
                "cleaned_sections": 0,
                "header_change": 0,
                "footer_change": 0,
            }
            return result

        # 用 python-docx 打开
        original = self._safe_open(original_path)
        cleaned = self._safe_open(cleaned_path)

        if original is None or cleaned is None:
            return result

        try:
            result["open_success"] = True

            # 结构检查
            structure = self._check_structure(original, cleaned)
            result["structure"] = structure

            # 内容检查
            content = self._check_content(original, cleaned)
            result["content"] = content

        finally:
            if original:
                original.close()
            if cleaned:
                cleaned.close()

        return result

    @staticmethod
    def _safe_open(path: str) -> Optional[Document]:
        """安全打开 DOCX，失败返回 None。"""
        try:
            return Document(path)
        except PackageNotFoundError:
            logger.error("DOCX 包未找到: %s", path)
            return None
        except Exception as e:
            logger.error("DOCX 打开失败: %s, 错误: %s", path, e)
            return None

    @staticmethod
    def _check_zip_structure(path: str) -> bool:
        """检查 DOCX 的 ZIP 结构完整性。"""
        try:
            with zipfile.ZipFile(path, "r") as zf:
                names = set(zf.namelist())
                for required in REQUIRED_XML_FILES:
                    if required not in names:
                        logger.error("DOCX 缺少必需文件: %s", required)
                        return False
                # 检查 _rels 目录存在
                has_rels = any(n.startswith("word/_rels/") for n in names)
                if not has_rels:
                    logger.error("DOCX 缺少 word/_rels 目录")
                    return False
            return True
        except Exception as e:
            logger.error("DOCX ZIP 检查失败: %s", e)
            return False

    @staticmethod
    def _check_structure(
        original: Document,
        cleaned: Document,
    ) -> Dict[str, object]:
        """检查文档结构变化。"""
        orig_sections = len(original.sections)
        clean_sections = len(cleaned.sections)

        # Header/Footer 数量
        orig_headers = sum(
            1 for s in original.sections
            if s.header and s.header.paragraphs
            and any(p.text.strip() for p in s.header.paragraphs)
        )
        clean_headers = sum(
            1 for s in cleaned.sections
            if s.header and s.header.paragraphs
            and any(p.text.strip() for p in s.header.paragraphs)
        )

        orig_footers = sum(
            1 for s in original.sections
            if s.footer and s.footer.paragraphs
            and any(p.text.strip() for p in s.footer.paragraphs)
        )
        clean_footers = sum(
            1 for s in cleaned.sections
            if s.footer and s.footer.paragraphs
            and any(p.text.strip() for p in s.footer.paragraphs)
        )

        return {
            "section_match": orig_sections == clean_sections,
            "original_sections": orig_sections,
            "cleaned_sections": clean_sections,
            "header_change": orig_headers - clean_headers,
            "footer_change": orig_footers - clean_footers,
            "original_headers": orig_headers,
            "cleaned_headers": clean_headers,
            "original_footers": orig_footers,
            "cleaned_footers": clean_footers,
        }

    @staticmethod
    def _check_content(
        original: Document,
        cleaned: Document,
    ) -> Dict[str, object]:
        """检查文档内容变化。"""
        orig_paras = len(original.paragraphs)
        clean_paras = len(cleaned.paragraphs)

        orig_text = " ".join(p.text for p in original.paragraphs)
        clean_text = " ".join(p.text for p in cleaned.paragraphs)

        orig_len = len(orig_text)
        clean_len = len(clean_text)

        text_loss_rate = 0.0
        if orig_len > 0:
            text_loss_rate = round((orig_len - clean_len) / orig_len, 4)

        return {
            "total_chars_before": orig_len,
            "total_chars_after": clean_len,
            "paragraphs_before": orig_paras,
            "paragraphs_after": clean_paras,
            "text_loss_rate": text_loss_rate,
            "text_changed": orig_text != clean_text,
        }
