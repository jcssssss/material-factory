"""ShapeDetector 测试。"""

from __future__ import annotations

import os
import tempfile
import zipfile
from xml.etree import ElementTree as ET

import pytest
from docx import Document

from detector.shape_detector import ShapeDetector

NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS_V = "urn:schemas-microsoft-com:vml"
NS_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def _create_docx_with_shape(path: str) -> None:
    """创建包含 VML Shape 的 DOCX。"""
    doc = Document()
    doc.add_paragraph("Body content")
    doc.save(path)

    # 在 header1.xml 中注入 VML Shape
    vml_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:hdr xmlns:w="{NS_W}" xmlns:v="{NS_V}" xmlns:r="{NS_R}">
  <w:p>
    <w:r>
      <w:pict>
        <v:shape id="Watermark" style="position:absolute;rotation:315" fillcolor="#CCCCCC"/>
      </w:pict>
    </w:r>
  </w:p>
  <w:p>
    <w:r><w:t>Normal header text</w:t></w:r>
  </w:p>
</w:hdr>"""

    import shutil
    fd, tmp_path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)

    with zipfile.ZipFile(path, "r") as zin:
        with zipfile.ZipFile(tmp_path, "w") as zout:
            has_header = False
            for item in zin.infolist():
                if item.filename == "word/header1.xml":
                    zout.writestr(item, vml_xml)
                    has_header = True
                else:
                    zout.writestr(item, zin.read(item.filename))
            if not has_header:
                zout.writestr("word/header1.xml", vml_xml)
                # 更新 [Content_Types].xml
                if "[Content_Types].xml" not in [n.filename for n in zin.infolist()]:
                    types_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/header1.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml"/>
</Types>"""
                    zout.writestr("[Content_Types].xml", types_xml)

    shutil.move(tmp_path, path)


def _create_docx_with_textbox(path: str) -> None:
    """创建包含 TextBox 的 DOCX。"""
    doc = Document()
    doc.add_paragraph("Body")
    doc.save(path)

    textbox_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:hdr xmlns:w="{NS_W}" xmlns:v="{NS_V}">
  <w:p>
    <w:r>
      <w:pict>
        <v:textbox>
          <w:txbxContent>
            <w:p><w:r><w:t>内部资料</w:t></w:r></w:p>
          </w:txbxContent>
        </v:textbox>
      </w:pict>
    </w:r>
  </w:p>
</w:hdr>"""

    import shutil
    fd, tmp_path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)

    with zipfile.ZipFile(path, "r") as zin:
        files = {item.filename: zin.read(item) for item in zin.infolist()}
    files["word/header1.xml"] = textbox_xml.encode("utf-8")
    with zipfile.ZipFile(tmp_path, "w") as zout:
        for filename, data in files.items():
            zout.writestr(filename, data)
    shutil.move(tmp_path, path)


class TestShapeDetector:
    """Shape 检测测试。"""

    def setup_method(self) -> None:
        self.detector = ShapeDetector()

    def test_detect_vml_shape(self) -> None:
        """应检测到 VML Shape 对象。"""
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        _create_docx_with_shape(path)

        try:
            objects = self.detector.detect(path)
            shapes = [o for o in objects if o.object_type == "shape"]
            assert len(shapes) >= 1
            assert shapes[0].confidence >= 0.7
        finally:
            os.unlink(path)

    def test_detect_textbox(self) -> None:
        """应检测到 TextBox 对象。"""
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        _create_docx_with_textbox(path)

        try:
            objects = self.detector.detect(path)
            textboxes = [o for o in objects if o.object_type == "textbox"]
            assert len(textboxes) >= 1
        finally:
            os.unlink(path)

    def test_normal_doc_no_shapes(self) -> None:
        """无 Shape 的文档应无检测结果。"""
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        doc = Document()
        doc.add_paragraph("Normal")
        doc.save(path)

        try:
            objects = self.detector.detect(path)
            shapes = [o for o in objects if o.object_type == "shape"]
            assert len(shapes) == 0
        finally:
            os.unlink(path)
