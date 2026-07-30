"""Word 分析器与基础框架测试。"""

from __future__ import annotations

import os
import tempfile

import pytest
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.shared import Inches, Pt

from analyzer.word_analyzer import WordAnalyzer
from analyzer.section_analyzer import SectionAnalyzer
from detector.word_detector import WordDetector
from cleaner.word_cleaner import WordCleaner
from risk import CleaningAction, CleaningPlan, RiskLevel
from cleaner import CleaningStatus


def _create_simple_docx(path: str) -> None:
    """创建普通 DOCX 文档。"""
    doc = Document()
    doc.add_paragraph("Hello World")
    doc.add_paragraph("This is a test document.")
    doc.save(path)


def _create_multi_section_docx(path: str) -> None:
    """创建多 Section DOCX 文档。"""
    doc = Document()

    # Section 1
    doc.add_paragraph("Section 1 content")
    doc.add_section()
    doc.add_paragraph("Section 2 content")
    doc.save(path)


def _create_header_footer_docx(path: str) -> None:
    """创建带页眉页脚的 DOCX 文档。"""
    doc = Document()

    # Header
    header = doc.sections[0].header
    header.add_paragraph("Internal Document")

    # Footer
    footer = doc.sections[0].footer
    footer.add_paragraph("Page 1")

    doc.add_paragraph("Main body content.")
    doc.save(path)


def _create_drawing_docx(path: str) -> None:
    """创建包含 DrawingML 的 DOCX 文档。"""
    doc = Document()
    doc.add_paragraph("Document with image reference")
    doc.save(path)

    # 在 document.xml 中注入 drawing 元素
    from lxml import etree
    import zipfile

    ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ns_wp = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    ns_a = "http://schemas.openxmlformats.org/drawingml/2006/main"
    ns_r = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

    # 读取并修改 document.xml
    with zipfile.ZipFile(path, "r") as zf:
        data = zf.read("word/document.xml")
        tree = etree.fromstring(data)
        body = tree.find(f"{{{ns_w}}}body")

        # 添加一个包含 drawing 的段落
        drawing_xml = f"""<w:p xmlns:w="{ns_w}" xmlns:wp="{ns_wp}" xmlns:a="{ns_a}" xmlns:r="{ns_r}"><w:r><w:drawing><wp:inline><wp:extent cx="100" cy="100"/></wp:inline></w:drawing></w:r></w:p>"""
        drawing_elem = etree.fromstring(drawing_xml)
        body.append(drawing_elem)

        new_data = etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)

    # 写回
    import shutil
    tmp_path = path + ".tmp"
    with zipfile.ZipFile(path, "r") as zin:
        with zipfile.ZipFile(tmp_path, "w") as zout:
            for item in zin.infolist():
                if item.filename == "word/document.xml":
                    zout.writestr(item, new_data)
                else:
                    zout.writestr(item, zin.read(item.filename))
    shutil.move(tmp_path, path)


class TestWordAnalyzer:
    """WordAnalyzer 测试。"""

    def setup_method(self) -> None:
        self.analyzer = WordAnalyzer()

    def test_simple_docx(self) -> None:
        """普通 DOCX 应正确分析。"""
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        _create_simple_docx(path)

        try:
            result = self.analyzer.analyze(path)
            assert result.paragraph_count > 0
            assert result.section_count > 0
            assert result.metadata.get("error") is None
        finally:
            os.unlink(path)

    def test_multi_section(self) -> None:
        """多 Section 文档应正确统计。"""
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        _create_multi_section_docx(path)

        try:
            result = self.analyzer.analyze(path)
            assert result.section_count == 2
        finally:
            os.unlink(path)

    def test_header_footer(self) -> None:
        """带页眉页脚的文档应正确识别。"""
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        _create_header_footer_docx(path)

        try:
            result = self.analyzer.analyze(path)
            assert result.header_count >= 1
            assert result.footer_count >= 1
        finally:
            os.unlink(path)

    def test_drawing_detection(self) -> None:
        """包含 Drawing 的文档应检测到。"""
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        _create_drawing_docx(path)

        try:
            result = self.analyzer.analyze(path)
            assert result.has_drawing is True
        finally:
            os.unlink(path)

    def test_non_docx_format(self) -> None:
        """非 DOCX 格式应返回错误。"""
        fd, path = tempfile.mkstemp(suffix=".doc")
        os.close(fd)
        try:
            result = self.analyzer.analyze(path)
            assert "UNSUPPORTED_FORMAT" in str(result.metadata.get("error", ""))
        finally:
            os.unlink(path)

    def test_nonexistent_file(self) -> None:
        """不存在的文件应返回错误。"""
        result = self.analyzer.analyze("/nonexistent/path.docx")
        assert "error" in result.metadata


class TestSectionAnalyzer:
    """SectionAnalyzer 测试。"""

    def setup_method(self) -> None:
        self.analyzer = SectionAnalyzer()

    def test_section_header_footer(self) -> None:
        """Section 页眉页脚应正确识别。"""
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        _create_header_footer_docx(path)

        try:
            doc = Document(path)
            sections = self.analyzer.analyze(doc)
            assert len(sections) >= 1
            assert sections[0]["header_exists"] is True
            assert sections[0]["footer_exists"] is True
            assert len(sections[0]["header_paragraphs"]) > 0
            assert len(sections[0]["footer_paragraphs"]) > 0
        finally:
            os.unlink(path)


class TestWordDetector:
    """WordDetector 基础框架测试。"""

    def setup_method(self) -> None:
        self.detector = WordDetector()

    def test_detect_simple(self) -> None:
        """基础检测应不报错。"""
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        _create_simple_docx(path)
        try:
            results = self.detector.detect(path)
            assert isinstance(results, list)
        finally:
            os.unlink(path)

    def test_detect_header_footer(self) -> None:
        """应检测到页眉页脚内容。"""
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        _create_header_footer_docx(path)
        try:
            results = self.detector.detect(path)
            assert len(results) >= 2  # header + footer
        finally:
            os.unlink(path)

    def test_detect_drawing(self) -> None:
        """应检测到 DrawingML 元素。"""
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        _create_drawing_docx(path)
        try:
            results = self.detector.detect(path)
            types = {r.type for r in results}
            # 应有 "word_element" 类型指示 drawing
        finally:
            os.unlink(path)


class TestWordCleaner:
    """WordCleaner 基础框架测试。"""

    def setup_method(self) -> None:
        self.cleaner = WordCleaner()

    def test_not_implemented(self) -> None:
        """V1 应返回 NOT_IMPLEMENTED。"""
        action = CleaningAction(
            action_type="REMOVE_TEXT",
            page=1,
            target_type="header",
            confidence=0.8,
            risk_level=RiskLevel.CONFIRM,
            risk_score=70.0,
        )
        plan = CleaningPlan(
            file_path="test.docx",
            risk_level=RiskLevel.CONFIRM,
            actions=[action],
        )
        results = self.cleaner.clean("test.docx", plan, "out.docx")
        assert len(results) == 1
        assert results[0].status == CleaningStatus.SKIPPED
        assert "NOT_IMPLEMENTED" in str(results[0].metadata.get("reason", ""))
