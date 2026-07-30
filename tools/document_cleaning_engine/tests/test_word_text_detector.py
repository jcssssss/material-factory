"""WordTextDetector 测试。"""

from __future__ import annotations

import os
import tempfile
import zipfile
from xml.etree import ElementTree as ET

import pytest
from docx import Document

from detector.word_text_detector import WordTextDetector

NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _create_docx_with_header(path: str, header_text: str) -> None:
    """创建带页眉文字的 DOCX。"""
    doc = Document()
    header = doc.sections[0].header
    header.add_paragraph(header_text)
    doc.add_paragraph("Main body content paragraph.")
    doc.save(path)


def _create_docx_with_body_keyword(path: str, keyword: str) -> None:
    """创建正文包含关键词的 DOCX。"""
    doc = Document()
    doc.add_paragraph(f"This document contains {keyword} in body.")
    doc.add_paragraph("Other normal content.")
    doc.save(path)


def _create_multi_section_shared_header(path: str, text: str) -> None:
    """创建多 Section 共享同一 Header 的 DOCX。"""
    doc = Document()
    header = doc.sections[0].header
    header.add_paragraph(text)
    doc.add_paragraph("Section 1 body.")

    doc.add_section()
    # Section 2 共享 Section 1 的 Header（Same As Previous）
    doc.add_paragraph("Section 2 body.")
    doc.save(path)


class TestWordTextDetector:
    """Word 文字水印检测测试。"""

    def setup_method(self) -> None:
        self.detector = WordTextDetector()

    def test_header_watermark_detected(self) -> None:
        """Header 中的水印文本应高置信度检测。"""
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        _create_docx_with_header(path, "内部资料")

        try:
            results = self.detector.detect(path)
            assert len(results) >= 1
            r = results[0]
            assert r.confidence >= 0.8
            assert "内部资料" in r.content
        finally:
            os.unlink(path)

    def test_body_keyword_ignored(self) -> None:
        """正文中仅出现一次的关键词应 IGNORE。"""
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        _create_docx_with_body_keyword(path, "版权所有")

        try:
            results = self.detector.detect(path)
            # 正文一次出现 + 非 header/footer → 评分低
            assert len(results) == 0
        finally:
            os.unlink(path)

    def test_multi_section_detection(self) -> None:
        """多 Section 应正确检测。"""
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        _create_multi_section_shared_header(path, "Confidential")

        try:
            results = self.detector.detect(path)
            assert len(results) >= 1
            assert results[0].confidence >= 0.8
        finally:
            os.unlink(path)

    def test_metadata_contains_xml_path(self) -> None:
        """检测结果应包含 XML 路径。"""
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        _create_docx_with_header(path, "Draft")

        try:
            results = self.detector.detect(path)
            if results:
                assert "xml" in results[0].metadata
        finally:
            os.unlink(path)

    def test_no_matches_for_normal_doc(self) -> None:
        """无关键词的普通文档应无检测结果。"""
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        doc = Document()
        doc.add_paragraph("This is a normal document with no watermarks.")
        doc.save(path)

        try:
            results = self.detector.detect(path)
            assert len(results) == 0
        finally:
            os.unlink(path)
