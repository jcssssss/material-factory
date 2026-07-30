"""WordObjectCleaner 测试。"""

from __future__ import annotations

import os
import tempfile
import zipfile
from xml.etree import ElementTree as ET

import pytest
from docx import Document

from cleaner.word_object_cleaner import WordObjectCleaner
from cleaner.shape_cleaner import ShapeCleaner
from cleaner.drawing_cleaner import DrawingCleaner
from risk import CleaningAction, CleaningPlan, RiskLevel
from cleaner import CleaningStatus

NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS_V = "urn:schemas-microsoft-com:vml"
NS_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def _create_docx_with_shape(path: str) -> None:
    """创建含 Shape 的 DOCX。"""
    doc = Document()
    doc.add_paragraph("Body content")
    doc.save(path)

    vml_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:hdr xmlns:w="{NS_W}" xmlns:v="{NS_V}" xmlns:r="{NS_R}">
  <w:p>
    <w:r>
      <w:pict>
        <v:shape id="Watermark" style="position:absolute;rotation:315"/>
      </w:pict>
    </w:r>
  </w:p>
  <w:p>
    <w:r><w:t>Normal header text</w:t></w:r>
  </w:p>
</w:hdr>"""

    import shutil
    fd, tmp_path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)

    with zipfile.ZipFile(path, "r") as zin:
        files = {item.filename: zin.read(item) for item in zin.infolist()}
        files["word/header1.xml"] = vml_xml.encode("utf-8")

    with zipfile.ZipFile(tmp_path, "w") as zout:
        for filename, data in files.items():
            zout.writestr(filename, data)
    shutil.move(tmp_path, path)


def _create_docx_with_drawing(path: str) -> None:
    """创建含 DrawingML 的 DOCX。"""
    ET.register_namespace("w", NS_W)

    doc = Document()
    doc.add_paragraph("Body")
    doc.save(path)

    import shutil
    drawing_xml = f"""<w:p xmlns:w="{NS_W}">
  <w:r>
    <w:drawing>
      <wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" allowOverlap="1">
        <wp:extent cx="914400" cy="914400"/>
        <wp:docPr id="1" name="Picture 1"/>
      </wp:anchor>
    </w:drawing>
  </w:r>
</w:p>"""

    with zipfile.ZipFile(path, "r") as zin:
        files = {item.filename: zin.read(item) for item in zin.infolist()}

    raw = files.get("word/document.xml", b"")
    root = ET.fromstring(raw)
    body = root.find(f"{{{NS_W}}}body")
    body.append(ET.fromstring(drawing_xml))
    files["word/document.xml"] = ET.tostring(root, xml_declaration=True, encoding="UTF-8")

    fd, tmp_path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    with zipfile.ZipFile(tmp_path, "w") as zout:
        for filename, data in files.items():
            zout.writestr(filename, data)
    shutil.move(tmp_path, path)


def _check_tag_exists(docx_path: str, tag_local: str) -> bool:
    """检查 DOCX 中是否包含指定标签。"""
    try:
        with zipfile.ZipFile(docx_path, "r") as zf:
            for name in zf.namelist():
                if not name.endswith(".xml"):
                    continue
                content = zf.read(name)
                if tag_local.encode() in content:
                    return True
    except Exception:
        pass
    return False


class TestShapeCleaner:
    """Shape 清理测试。"""

    def setup_method(self) -> None:
        self.cleaner = ShapeCleaner()

    def test_delete_shape(self) -> None:
        """应删除 v:shape 节点。"""
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        _create_docx_with_shape(path)

        try:
            assert _check_tag_exists(path, "v:shape")
            result = self.cleaner.delete_shape_in_xml(path)
            assert result is True
            assert not _check_tag_exists(path, "v:shape")
        finally:
            os.unlink(path)


class TestDrawingCleaner:
    """Drawing 清理测试。"""

    def setup_method(self) -> None:
        self.cleaner = DrawingCleaner()

    def test_delete_drawing(self) -> None:
        """应删除 w:drawing 节点。"""
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        _create_docx_with_drawing(path)

        try:
            assert _check_tag_exists(path, "w:drawing")
            result = self.cleaner.delete_drawing_in_xml(path)
            assert result is True
            assert not _check_tag_exists(path, "w:drawing")
        finally:
            os.unlink(path)


class TestWordObjectCleaner:
    """WordObjectCleaner 集成测试。"""

    def setup_method(self) -> None:
        self.cleaner = WordObjectCleaner()

    def test_clean_shape_action(self) -> None:
        """Shape 删除 Action 应执行成功。"""
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        _create_docx_with_shape(path)
        fd2, output_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd2)

        try:
            action = CleaningAction(
                action_type="REMOVE_SHAPE",
                page=0, target_type="word_object",
                confidence=0.9, risk_level=RiskLevel.AUTO, risk_score=85.0,
                metadata={"object_type": "shape", "xml_file": "word/header1.xml"},
            )
            plan = CleaningPlan(
                file_path=path, risk_level=RiskLevel.AUTO, actions=[action],
            )
            results = self.cleaner.clean(path, plan, output_path)
            assert len(results) == 1
            assert results[0].status == CleaningStatus.SUCCESS
            assert not _check_tag_exists(output_path, "v:shape")
        finally:
            if os.path.exists(path):
                os.unlink(path)
            if os.path.exists(output_path):
                os.unlink(output_path)

    def test_clean_drawing_action(self) -> None:
        """Drawing 删除 Action 应执行成功。"""
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        _create_docx_with_drawing(path)
        fd2, output_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd2)

        try:
            action = CleaningAction(
                action_type="REMOVE_DRAWING",
                page=0, target_type="word_object",
                confidence=0.9, risk_level=RiskLevel.AUTO, risk_score=85.0,
                metadata={"object_type": "drawing", "xml_file": "word/document.xml"},
            )
            plan = CleaningPlan(
                file_path=path, risk_level=RiskLevel.AUTO, actions=[action],
            )
            results = self.cleaner.clean(path, plan, output_path)
            assert len(results) == 1
            assert results[0].status == CleaningStatus.SUCCESS
            assert not _check_tag_exists(output_path, "w:drawing")
        finally:
            if os.path.exists(path):
                os.unlink(path)
            if os.path.exists(output_path):
                os.unlink(output_path)

    def test_unknown_object_type(self) -> None:
        """未知对象类型应返回 FAILED。"""
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        doc = Document()
        doc.add_paragraph("Test")
        doc.save(path)
        fd2, output_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd2)

        try:
            action = CleaningAction(
                action_type="REMOVE_UNKNOWN",
                page=0, target_type="word_object",
                confidence=0.5, risk_level=RiskLevel.AUTO, risk_score=50.0,
                metadata={"object_type": "unknown"},
            )
            plan = CleaningPlan(
                file_path=path, risk_level=RiskLevel.AUTO, actions=[action],
            )
            results = self.cleaner.clean(path, plan, output_path)
            assert len(results) == 1
            assert results[0].status == CleaningStatus.FAILED
        finally:
            if os.path.exists(path):
                os.unlink(path)
