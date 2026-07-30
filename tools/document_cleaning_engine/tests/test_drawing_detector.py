"""DrawingDetector 测试。"""

from __future__ import annotations

import os
import tempfile
import zipfile
from xml.etree import ElementTree as ET

import pytest
from docx import Document

from detector.drawing_detector import DrawingDetector

NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS_WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
NS_A = "http://schemas.openxmlformats.org/drawingml/2006/main"
NS_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
NS_PIC = "http://schemas.openxmlformats.org/drawingml/2006/picture"


def _create_docx_with_drawing(path: str) -> None:
    """创建包含 DrawingML 的 DOCX。"""
    # 注册命名空间确保序列化正确
    ET.register_namespace("w", NS_W)
    ET.register_namespace("wp", NS_WP)
    ET.register_namespace("a", NS_A)
    ET.register_namespace("r", NS_R)

    doc = Document()
    doc.add_paragraph("Body")
    doc.save(path)

    import shutil
    fd, tmp_path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)

    with zipfile.ZipFile(path, "r") as zin:
        files = {item.filename: zin.read(item) for item in zin.infolist()}

    raw = files.get("word/document.xml", b"")
    root = ET.fromstring(raw)
    body = root.find(f"{{{NS_W}}}body")

    drawing_xml = f"""<w:p xmlns:w="{NS_W}">
  <w:r>
    <w:drawing>
      <wp:anchor xmlns:wp="{NS_WP}" allowOverlap="1">
        <wp:extent cx="914400" cy="914400"/>
        <wp:docPr id="1" name="Picture 1"/>
      </wp:anchor>
    </w:drawing>
  </w:r>
</w:p>"""

    body.append(ET.fromstring(drawing_xml))
    new_raw = ET.tostring(root, xml_declaration=True, encoding="UTF-8")
    files["word/document.xml"] = new_raw

    with zipfile.ZipFile(tmp_path, "w") as zout:
        for filename, data in files.items():
            zout.writestr(filename, data)
    shutil.move(tmp_path, path)


class TestDrawingDetector:
    """DrawingML 检测测试。"""

    def setup_method(self) -> None:
        self.detector = DrawingDetector()

    def test_detect_anchor_drawing(self) -> None:
        """应检测到浮动 DrawingML 对象。"""
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        _create_docx_with_drawing(path)

        try:
            objects = self.detector.detect(path)
            drawings = [o for o in objects if "drawing" in o.object_type]
            assert len(drawings) >= 1
        finally:
            os.unlink(path)

    def test_normal_doc_no_drawings(self) -> None:
        """无 Drawing 的文档应无检测结果。"""
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        doc = Document()
        doc.add_paragraph("Normal")
        doc.save(path)

        try:
            objects = self.detector.detect(path)
            drawings = [o for o in objects if "drawing" in o.object_type]
            assert len(drawings) == 0
        finally:
            os.unlink(path)
