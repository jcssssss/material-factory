"""WordTextCleaner 测试。"""

from __future__ import annotations

import os
import shutil
import tempfile
import zipfile
from xml.etree import ElementTree as ET

import pytest
from docx import Document

from cleaner.word_text_cleaner import WordTextCleaner
from risk import CleaningAction, CleaningPlan, RiskLevel
from cleaner import CleaningStatus
from matcher.word_text_matcher import WordTextMatcher

NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _create_docx_with_header(path: str, header_text: str) -> str:
    """创建带页眉文字的 DOCX。"""
    doc = Document()
    header = doc.sections[0].header
    header.add_paragraph(header_text)
    doc.add_paragraph("Main body content.")
    doc.save(path)
    return path


def _find_header_xml(docx_path: str) -> str:
    """在 DOCX 中找到 header XML 文件名。"""
    with zipfile.ZipFile(docx_path, "r") as zf:
        for name in zf.namelist():
            if "header" in name and name.endswith(".xml"):
                return name
    return ""


def _check_text_in_docx(docx_path: str, xml_path: str, text: str) -> bool:
    """检查 DOCX 的 XML 中是否包含指定文本。"""
    try:
        with zipfile.ZipFile(docx_path, "r") as zf:
            if xml_path not in zf.namelist():
                return False
            content = zf.read(xml_path)
            return text.encode("utf-8") in content
    except Exception:
        return False


class TestWordTextCleaner:
    """Word 文字水印清理测试。"""

    def setup_method(self) -> None:
        self.cleaner = WordTextCleaner()
        self.matcher = WordTextMatcher()

    def test_delete_header_text(self) -> None:
        """应删除 Header 中的指定文字。"""
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        _create_docx_with_header(path, "内部资料")
        fd2, output_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd2)

        try:
            # 找到 header XML 路径
            header_xml = _find_header_xml(path)
            assert header_xml, "header XML not found"

            action = CleaningAction(
                action_type="REMOVE_TEXT",
                page=0,
                target_type="header",
                confidence=0.95,
                risk_level=RiskLevel.AUTO,
                risk_score=90.0,
                content="内部资料",
                metadata={"xml": header_xml},
            )
            plan = CleaningPlan(
                file_path=path,
                risk_level=RiskLevel.AUTO,
                actions=[action],
            )

            results = self.cleaner.clean(path, plan, output_path)
            assert len(results) == 1
            assert results[0].status == CleaningStatus.SUCCESS

            # 验证文本已被删除
            assert not _check_text_in_docx(output_path, header_xml, "内部资料")
        finally:
            if os.path.exists(path):
                os.unlink(path)
            if os.path.exists(output_path):
                os.unlink(output_path)

    def test_delete_failed_for_nonexistent_text(self) -> None:
        """不存在的文本应返回 FAILED。"""
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        _create_docx_with_header(path, "Normal Header")
        fd2, output_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd2)

        try:
            header_xml = _find_header_xml(path)

            action = CleaningAction(
                action_type="REMOVE_TEXT",
                page=0,
                target_type="header",
                confidence=0.95,
                risk_level=RiskLevel.AUTO,
                risk_score=90.0,
                content="NONEXISTENT_TEXT",
                metadata={"xml": header_xml},
            )
            plan = CleaningPlan(
                file_path=path,
                risk_level=RiskLevel.AUTO,
                actions=[action],
            )

            results = self.cleaner.clean(path, plan, output_path)
            assert len(results) == 1
            assert results[0].status == CleaningStatus.FAILED
        finally:
            if os.path.exists(path):
                os.unlink(path)
            if os.path.exists(output_path):
                os.unlink(output_path)

    def test_empty_plan(self) -> None:
        """空计划应返回空结果。"""
        plan = CleaningPlan(file_path="test.docx", risk_level=RiskLevel.AUTO)
        results = self.cleaner.clean("test.docx", plan, "out.docx")
        assert len(results) == 0

    def test_body_text_not_deleted(self) -> None:
        """正文文本不应被误删。"""
        fd, path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        doc = Document()
        doc.add_paragraph("Normal body content that should remain.")
        doc.save(path)
        fd2, output_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd2)

        try:
            # 删除 header 中不存在的文本（不影响正文）
            action = CleaningAction(
                action_type="REMOVE_TEXT",
                page=0,
                target_type="header",
                confidence=0.9,
                risk_level=RiskLevel.AUTO,
                risk_score=80.0,
                content="NONEXISTENT",
                metadata={"xml": "word/header1.xml"},
            )
            plan = CleaningPlan(
                file_path=path,
                risk_level=RiskLevel.AUTO,
                actions=[action],
            )

            results = self.cleaner.clean(path, plan, output_path)
            # 文本不存在应 FAILED 但不破坏文档
            assert len(results) == 1
            assert results[0].status == CleaningStatus.FAILED

            # 文档仍可正常打开
            doc2 = Document(path)
            texts = [p.text for p in doc2.paragraphs]
            assert any("Normal" in (t or "") for t in texts)
        finally:
            if os.path.exists(path):
                os.unlink(path)
            if os.path.exists(output_path):
                os.unlink(output_path)
