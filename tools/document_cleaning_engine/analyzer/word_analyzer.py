"""Word 文档分析器。

分析 DOCX 文件结构，输出 WordDocument 分析结果。
"""

from __future__ import annotations

import logging
import zipfile
from typing import Optional
from xml.etree import ElementTree as ET

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from models import WordDocument

from .section_analyzer import SectionAnalyzer

logger = logging.getLogger(__name__)

# OOXML 命名空间
NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS_WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
NS_A = "http://schemas.openxmlformats.org/drawingml/2006/main"
NS_V = "urn:schemas-microsoft-com:vml"
NS_WPC = "http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"
NS_MC = "http://schemas.openxmlformats.org/markup-compatibility/2006"
NS_W14 = "http://schemas.microsoft.com/office/word/2010/wordml"


class WordAnalyzer:
    """Word 文档分析器。

    使用 python-docx 进行高层结构分析，zipfile + lxml 进行底层 XML 扫描。
    """

    def __init__(self) -> None:
        self._section_analyzer = SectionAnalyzer()

    def analyze(self, docx_path: str) -> WordDocument:
        """分析 DOCX 文件。

        Args:
            docx_path: DOCX 文件路径。

        Returns:
            WordDocument 分析结果。
        """
        # 文件后缀检查
        if not docx_path.lower().endswith(".docx"):
            return WordDocument(
                file_path=docx_path,
                paragraph_count=0,
                section_count=0,
                header_count=0,
                footer_count=0,
                has_shapes=False,
                has_drawing=False,
                metadata={"error": "UNSUPPORTED_FORMAT", "format": "non-docx"},
            )

        # 尝试打开
        try:
            doc = Document(docx_path)
        except PackageNotFoundError:
            return WordDocument(
                file_path=docx_path,
                paragraph_count=0,
                section_count=0,
                header_count=0,
                footer_count=0,
                has_shapes=False,
                has_drawing=False,
                metadata={"error": "INVALID_DOCX"},
            )
        except Exception as e:
            error_msg = str(e).lower()
            if "encrypted" in error_msg or "password" in error_msg:
                return WordDocument(
                    file_path=docx_path,
                    paragraph_count=0,
                    section_count=0,
                    header_count=0,
                    footer_count=0,
                    has_shapes=False,
                    has_drawing=False,
                    metadata={"error": "ENCRYPTED_DOCUMENT"},
                )
            return WordDocument(
                file_path=docx_path,
                paragraph_count=0,
                section_count=0,
                header_count=0,
                footer_count=0,
                has_shapes=False,
                has_drawing=False,
                metadata={"error": f"OPEN_FAILED: {e}"},
            )

        try:
            # 基础统计
            paragraph_count = len(doc.paragraphs)

            # Section 分析
            sections = self._section_analyzer.analyze(doc)
            section_count = len(sections)

            header_count = sum(1 for s in sections if s["header_exists"])
            footer_count = sum(1 for s in sections if s["footer_exists"])

            # XML 扫描（via lxml/zipfile）
            has_shapes, has_drawing = self._scan_xml(docx_path)

            # 收集 Section 详情
            section_details = []
            for s in sections:
                section_details.append({
                    "index": s["section_index"],
                    "header_exists": s["header_exists"],
                    "footer_exists": s["footer_exists"],
                    "different_first_page": s["different_first_page"],
                    "header_texts": s["header_paragraphs"],
                    "footer_texts": s["footer_paragraphs"],
                })

            return WordDocument(
                file_path=docx_path,
                paragraph_count=paragraph_count,
                section_count=section_count,
                header_count=header_count,
                footer_count=footer_count,
                has_shapes=has_shapes,
                has_drawing=has_drawing,
                metadata={
                    "sections": section_details,
                    "error": None,
                },
            )
        except Exception as e:
            logger.error("Word analysis failed: %s", e)
            return WordDocument(
                file_path=docx_path,
                paragraph_count=0,
                section_count=0,
                header_count=0,
                footer_count=0,
                has_shapes=False,
                has_drawing=False,
                metadata={"error": f"ANALYSIS_FAILED: {e}"},
            )

    def _scan_xml(self, docx_path: str):
        """扫描 DOCX 的 XML 文件，检测 Shape 和 Drawing。

        Args:
            docx_path: DOCX 文件路径。

        Returns:
            (has_shapes, has_drawing) 布尔元组。
        """
        has_shapes = False
        has_drawing = False

        try:
            with zipfile.ZipFile(docx_path, "r") as zf:
                xml_files = [
                    n for n in zf.namelist()
                    if n.startswith("word/") and n.endswith(".xml")
                ]

                for xml_path in xml_files:
                    try:
                        content = zf.read(xml_path)
                        root = ET.fromstring(content)
                    except Exception:
                        continue

                    # 搜索 DrawingML drawing 元素
                    if not has_drawing:
                        has_drawing = self._search_drawing(root)

                    # 搜索 VML pict/shape 元素
                    if not has_shapes:
                        has_shapes = self._search_shapes(root)

                    if has_shapes and has_drawing:
                        break
        except Exception:
            pass

        return has_shapes, has_drawing

    def _search_drawing(self, root: ET.Element) -> bool:
        """搜索 XML 树中的 DrawingML 元素。"""
        # 直接搜索 w:drawing
        for elem in root.iter(f"{{{NS_W}}}drawing"):
            if elem is not None:
                return True
        return False

    def _search_shapes(self, root: ET.Element) -> bool:
        """搜索 XML 树中的 VML Shape 元素。"""
        # 搜索 w:pict（VML 图片包装）
        for elem in root.iter(f"{{{NS_W}}}pict"):
            if elem is not None:
                return True
        # 搜索 v:shape
        for elem in root.iter(f"{{{NS_V}}}shape"):
            if elem is not None:
                return True
        # 搜索 v:rect
        for elem in root.iter(f"{{{NS_V}}}rect"):
            if elem is not None:
                return True
        return False
